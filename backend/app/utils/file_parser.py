"""
ChronosIntel — Multi-format Document Parser
============================================
Supports: .txt, .md, .pdf, .docx, .csv, .json, .eml, .log

Each parser returns a `ParsedDocument` with:
  - text content
  - metadata (author, date, subject, etc.)
  - detected document type
"""

from __future__ import annotations

import csv
import email
import email.policy
import io
import json
import logging
import os
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from app.database.models import DocumentType
from app.utils.exceptions import DocumentProcessingError, UnsupportedFileTypeError

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Result of parsing a document file."""

    text: str
    document_type: DocumentType = DocumentType.DOCUMENT
    author: str | None = None
    document_date: datetime | None = None
    subject: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


class FileParser:
    """
    Parses various document formats into plain text with metadata.

    Usage::

        parser = FileParser()
        result = await parser.parse(file_path="/storage/uploads/email.eml")
    """

    EXTENSION_MAP: dict[str, str] = {
        ".txt": "_parse_text",
        ".md": "_parse_text",
        ".log": "_parse_text",
        ".pdf": "_parse_pdf",
        ".docx": "_parse_docx",
        ".doc": "_parse_docx",
        ".csv": "_parse_csv",
        ".json": "_parse_json",
        ".eml": "_parse_eml",
        ".msg": "_parse_eml",
    }

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a file and return its text content + metadata.

        Args:
            file_path: Absolute path to the file.

        Returns:
            ParsedDocument with extracted text and metadata.

        Raises:
            UnsupportedFileTypeError: If the file extension is not supported.
            DocumentProcessingError: If parsing fails.
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        filename = path.name

        if extension not in self.EXTENSION_MAP:
            raise UnsupportedFileTypeError(filename=filename, extension=extension)

        method_name = self.EXTENSION_MAP[extension]
        method = getattr(self, method_name)

        try:
            logger.debug("Parsing file: %s (type=%s)", filename, extension)
            result: ParsedDocument = method(path)
            logger.debug("Parsed %d chars from %s", len(result.text), filename)
            return result
        except (UnsupportedFileTypeError, DocumentProcessingError):
            raise
        except Exception as exc:
            raise DocumentProcessingError(
                filename=filename, reason=str(exc)
            ) from exc

    # ── Format Parsers ────────────────────────────────────────────────────────

    def _parse_text(self, path: Path) -> ParsedDocument:
        """Parse plain text / markdown / log files."""
        text = path.read_text(encoding="utf-8", errors="replace")
        doc_type = DocumentType.AUDIT_LOG if path.suffix == ".log" else DocumentType.DOCUMENT
        return ParsedDocument(text=text, document_type=doc_type)

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        """Extract text from PDF using pypdf."""
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise DocumentProcessingError(
                filename=path.name, reason="pypdf not installed. Run: pip install pypdf"
            ) from exc

        reader = PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)

        full_text = "\n\n".join(pages)
        metadata: dict[str, Any] = {}
        if reader.metadata:
            metadata["title"] = reader.metadata.get("/Title")
            metadata["author"] = reader.metadata.get("/Author")
            metadata["creation_date"] = str(reader.metadata.get("/CreationDate", ""))
            metadata["page_count"] = len(reader.pages)

        author = metadata.get("author")
        return ParsedDocument(
            text=full_text,
            document_type=DocumentType.DOCUMENT,
            author=author if isinstance(author, str) else None,
            metadata=metadata,
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        """Extract text from DOCX using python-docx."""
        try:
            import docx
        except ImportError as exc:
            raise DocumentProcessingError(
                filename=path.name,
                reason="python-docx not installed. Run: pip install python-docx",
            ) from exc

        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        # Extract core properties
        props = doc.core_properties
        author = props.author or None
        doc_date: datetime | None = None
        if props.created:
            doc_date = props.created.replace(tzinfo=timezone.utc) if props.created.tzinfo is None else props.created

        metadata = {
            "title": props.title,
            "author": author,
            "subject": props.subject,
            "paragraph_count": len(paragraphs),
        }
        return ParsedDocument(
            text=full_text,
            document_type=DocumentType.DOCUMENT,
            author=author,
            document_date=doc_date,
            metadata=metadata,
        )

    def _parse_csv(self, path: Path) -> ParsedDocument:
        """Convert CSV to readable text rows."""
        rows: list[str] = []
        with path.open(encoding="utf-8", errors="replace", newline="") as f:
            reader = csv.DictReader(f)
            fieldnames = reader.fieldnames or []
            rows.append(", ".join(fieldnames))
            for row in reader:
                rows.append(", ".join(str(v) for v in row.values()))

        return ParsedDocument(
            text="\n".join(rows),
            document_type=DocumentType.AUDIT_LOG,
            metadata={"row_count": len(rows) - 1, "columns": fieldnames},
        )

    def _parse_json(self, path: Path) -> ParsedDocument:
        """Parse JSON files — flatten to readable key-value text."""
        raw = json.loads(path.read_text(encoding="utf-8", errors="replace"))

        def flatten(obj: Any, prefix: str = "") -> list[str]:
            lines: list[str] = []
            if isinstance(obj, dict):
                for k, v in obj.items():
                    lines.extend(flatten(v, prefix=f"{prefix}{k}."))
            elif isinstance(obj, list):
                for i, item in enumerate(obj[:100]):  # cap at 100 items
                    lines.extend(flatten(item, prefix=f"{prefix}[{i}]."))
            else:
                lines.append(f"{prefix.rstrip('.')}: {obj}")
            return lines

        text = "\n".join(flatten(raw))
        return ParsedDocument(
            text=text,
            document_type=DocumentType.AUDIT_LOG,
            metadata={"raw_type": type(raw).__name__},
        )

    def _parse_eml(self, path: Path) -> ParsedDocument:
        """
        Parse .eml email files.
        Extracts From, To, Subject, Date and body text.
        """
        raw = path.read_bytes()
        msg = email.message_from_bytes(raw, policy=email.policy.default)

        subject = str(msg.get("Subject", ""))
        from_ = str(msg.get("From", ""))
        to_ = str(msg.get("To", ""))
        date_str = str(msg.get("Date", ""))

        # Parse date
        doc_date: datetime | None = None
        if date_str:
            try:
                from email.utils import parsedate_to_datetime
                doc_date = parsedate_to_datetime(date_str)
            except Exception:
                pass

        # Extract body
        body_parts: list[str] = []
        if msg.is_multipart():
            for part in msg.walk():
                ct = part.get_content_type()
                if ct == "text/plain":
                    payload = part.get_payload(decode=True)
                    if isinstance(payload, bytes):
                        body_parts.append(payload.decode("utf-8", errors="replace"))
        else:
            payload = msg.get_payload(decode=True)
            if isinstance(payload, bytes):
                body_parts.append(payload.decode("utf-8", errors="replace"))

        body = "\n".join(body_parts)
        full_text = (
            f"Subject: {subject}\n"
            f"From: {from_}\n"
            f"To: {to_}\n"
            f"Date: {date_str}\n\n"
            f"{body}"
        )

        # Extract author name from "From" header
        author_match = re.match(r"^(.+?)\s*<", from_)
        author = author_match.group(1).strip() if author_match else from_.split("<")[0].strip() or None

        return ParsedDocument(
            text=full_text,
            document_type=DocumentType.EMAIL,
            author=author or None,
            document_date=doc_date,
            subject=subject or None,
            metadata={
                "from": from_,
                "to": to_,
                "subject": subject,
                "date_raw": date_str,
            },
        )


# ── Module-level singleton ────────────────────────────────────────────────────
file_parser = FileParser()
